"""Markdown -> Word, for the reader who has to work on the document.

A ministry receives a report and then edits it: lifts a section into a
submission, comments in the margin, tracks changes through three offices. HTML
and PDF are read-only to that reader, so the document stops being usable at
exactly the point it starts being useful.

This is a **working document**, not a rendition of the print edition. Structure,
tables, diagrams and the source annex come across intact and land on real Word
styles, so an official can restyle to a house template without unpicking direct
formatting. The banners, gradients and measured page numbers do not: Word
paginates the document itself, and a colour ramp is not what survives being
pasted into someone else's template.

It is the pipeline's third emitter. The first two drifted apart within a day of
the second existing, so `structure()` reports what a docx carries and `verify`
compares it against the reading edition rather than trusting all three to agree.
"""
import html as ihtml
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

from . import assemble
from . import front as front_mod, markdown as md, palette, typst, xref

HEAD_RE = md.HEAD_RE
ATTR_RE = md.ATTR_RE
LIST_RE = md.LIST_RE
INLINE_CODE = re.compile(r'`([^`]+)`')
BOLD = re.compile(r'\*\*(.+?)\*\*')
ITALIC = re.compile(r'(?<!\*)\*([^*\n]+)\*(?!\*)')
LINK = re.compile(r'\[([^\]]+)\]\(([^)\s]+)\)')
WIDE = 6                      # columns from which a table gets a landscape section


def _colour(brand, token):
    """A palette token as a Word colour.

    This emitter used to keep its own four-entry copy of the defaults, of which
    it read two: `amber` and `muted` were declared, never used, and documented
    as reaching the Word edition. They do now.
    """
    return RGBColor(*palette.channels(palette.resolve(None, brand)[token]))


def _navy(brand):
    return _colour(brand, 'navy')


def _runs(paragraph, text):
    """Inline markdown onto Word runs. Links become their text plus the target,
    because a run carrying a relationship is not what survives a paste."""
    text = LINK.sub(lambda m: m.group(1) if m.group(2).startswith('./')
                    else '%s (%s)' % (m.group(1), m.group(2)), text)
    text = re.sub(r'<br\s*/?>', '\n', text)
    pos = 0
    for m in re.finditer(r'\*\*(.+?)\*\*|(?<!\*)\*([^*\n]+)\*(?!\*)|`([^`]+)`', text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold, italic, code = m.group(1), m.group(2), m.group(3)
        run = paragraph.add_run(bold or italic or code)
        run.bold, run.italic = bool(bold), bool(italic)
        if code:
            run.font.name = 'Consolas'
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _columns(section, n, space=425):
    """Set a section's column count. `w:cols` is already in every sectPr Word
    writes, so this sets attributes rather than appending an element out of
    schema order. 425 twips is 0.75cm, Word's own default gutter."""
    from docx.oxml.ns import qn
    cols = section._sectPr.find(qn('w:cols'))
    if cols is None:      # pragma: no cover - python-docx writes one every time
        from docx.oxml import OxmlElement
        cols = OxmlElement('w:cols')
        section._sectPr.append(cols)
    cols.set(qn('w:num'), str(n))
    cols.set(qn('w:space'), str(space))
    return section


def _line_numbers(doc):
    """Number every line, in every section.

    Word carries line numbering per section, and python-docx does not expose
    it, so the element goes in directly. Two details, both checked rather than
    assumed:

    `w:lnNumType` precedes `w:cols` in the schema's sequence, and appending it
    put it after `w:cols` and `w:docGrid`. That was harmless while nothing wrote
    `w:cols`; the column count does, so it is inserted in order instead.

    Every section, rather than the first: `add_section` clones the trailing
    section properties, so setting it early does in fact reach the sections a
    wide table opens later - measured, having expected otherwise. Relying on
    that means line numbering depends on the order two unrelated features run
    in, which is not a property worth keeping.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        marks = OxmlElement('w:lnNumType')
        marks.set(qn('w:countBy'), '1')
        marks.set(qn('w:restart'), 'continuous')
        cols = section._sectPr.find(qn('w:cols'))
        if cols is None:  # pragma: no cover - as above
            section._sectPr.append(marks)
        else:
            cols.addprevious(marks)


def _landscape(doc, on, columns=1):
    """Word changes page orientation per section, so a wide table gets its own.

    A wide table also leaves the columns behind: it needs the long edge of the
    paper at 8pt already, and half of that is not a smaller version of the
    problem. The body returns to its own column count after it.
    """
    section = doc.add_section()
    if on:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Mm(297), Mm(210)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin = section.right_margin = Mm(18)
    _columns(section, 1 if on else columns)
    return section


def _table(doc, rows, wide):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.add_row().cells
    for i, cell in enumerate(rows[0]):
        para = header[i].paragraphs[0]
        _runs(para, cell)
        for run in para.runs:
            run.bold = True
    for row in rows[2:]:
        cells = table.add_row().cells
        for i in range(len(rows[0])):
            para = cells[i].paragraphs[0]
            _runs(para, row[i] if i < len(row) else '')
            if wide:
                for run in para.runs:
                    run.font.size = Pt(8)
    return table


def take_caption(lines, pos, table):
    """Consume a `: text {#fig-x}` line after a block, as the other emitters do."""
    look = pos
    while look < len(lines) and not lines[look].strip():
        look += 1
    if look < len(lines):
        m = xref.CAPTION_RE.match(lines[look].strip())
        if m:
            return table.get(m.group(2)), look + 1
    return None, pos


def convert(doc, lines, figures, label, images, brand, part_banner=None,
            force_parts=False, table=None, columns=1):
    """Block pass. Mirrors the HTML emitter's structure decisions."""
    pos, n = 0, len(lines)
    landscape = False
    while pos < n:
        line, stripped = lines[pos], lines[pos].strip()
        if not stripped:
            pos += 1
            continue

        if stripped.startswith('```'):
            lang = stripped[3:].strip().lower()
            pos += 1
            buf = []
            while pos < n and not lines[pos].strip().startswith('```'):
                buf.append(lines[pos])
                pos += 1
            pos += 1
            if lang == 'mermaid':
                idx = len(figures)
                figures.append(idx)
                image = images.get(idx)
                if image and Path(image).exists():
                    doc.add_picture(str(image), width=Mm(160))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                entry, pos = take_caption(lines, pos, table or {})
                cap = doc.add_paragraph(xref.caption_of(entry) if entry
                                        else label % (idx + 1))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = _colour(brand, 'ink-soft')
            else:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(buf))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue

        if re.fullmatch(r'(-{3,}|\*{3,}|_{3,})', stripped):
            pos += 1
            continue

        m = HEAD_RE.match(stripped)
        if m:
            depth, text = len(m.group(1)), m.group(2)
            attrs = ''
            a = ATTR_RE.search(text)
            if a:
                attrs, text = a.group(1), ATTR_RE.sub('', text).strip()
            inferred = bool(part_banner and re.match(part_banner, text))
            is_part = ('part' in attrs) or force_parts or (inferred and 'no-part' not in attrs)
            if landscape:
                _landscape(doc, False, columns)
                landscape = False
            heading = doc.add_heading(level=min(depth, 4))
            _runs(heading, text)
            if is_part and depth == 2:
                for run in heading.runs:
                    run.font.color.rgb = _navy(brand)
            pos += 1
            continue

        expr, ident, after = xref.take_equation(lines, pos)
        if expr is not None:
            pos = after
            entry = (table or {}).get(ident) if ident else None
            # Word has its own equation editor; a plain centred line keeps the
            # expression readable and editable, which is what this edition is for
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(expr.strip())
            run.italic = True
            if entry:
                para.add_run('    (%d)' % entry['number'])
            continue

        if stripped.startswith('>'):
            buf = []
            while pos < n and lines[pos].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', lines[pos]))
                pos += 1
            first = buf[0] if buf else ''
            mark = re.match(r'^\[!(\w+)\]\s*(.*)$', first.strip())
            if mark:
                buf[0] = mark.group(2)
            para = doc.add_paragraph(style='Intense Quote')
            _runs(para, ' '.join(b.strip() for b in buf if b.strip()))
            # the type was matched and discarded here as well, so a warning read
            # as a note in the third edition too. Intense Quote owns the border;
            # the run colour is what this emitter can say, and it says the same
            # thing the rule down the edge says in the other two.
            for run in para.runs:
                run.font.color.rgb = _colour(brand, palette.variant(
                    mark.group(1) if mark else 'note')[0])
            continue

        if stripped.startswith('|') and pos + 1 < n and \
                re.match(r'^\|[\s:*|-]+\|?$', lines[pos + 1].strip()):
            rows = []
            while pos < n and lines[pos].lstrip().startswith('|'):
                rows.append([c.strip() for c in lines[pos].strip().strip('|').split('|')])
                pos += 1
            if len(rows) >= 2:
                wide = len(rows[0]) >= WIDE
                if wide and not landscape:
                    _landscape(doc, True, columns)
                    landscape = True
                elif landscape and not wide:
                    _landscape(doc, False, columns)
                    landscape = False
                _table(doc, rows, wide)
                entry, pos = take_caption(lines, pos, table or {})
                if entry:
                    cap = doc.add_paragraph(xref.caption_of(entry))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = _colour(brand, 'ink-soft')
                doc.add_paragraph()
            continue

        m = LIST_RE.match(line)
        if m:
            indent = len(m.group(1).expandtabs(4))
            ordered = m.group(2)[-1] == '.'
            style = 'List Number' if ordered else 'List Bullet'
            para = doc.add_paragraph(style=style)
            if indent >= 2:
                para.paragraph_format.left_indent = Mm(10 + 6 * (indent // 2))
            _runs(para, m.group(3))
            pos += 1
            continue

        buf = []
        while pos < n and lines[pos].strip() and \
                not lines[pos].strip().startswith(('```', '>', '|', '#')) and \
                not LIST_RE.match(lines[pos]):
            buf.append(lines[pos].strip())
            pos += 1
        _runs(doc.add_paragraph(), xref.substitute(' '.join(buf), table or {}))
    if landscape:
        _landscape(doc, False, columns)


def build(source, output, prof, svgs=None, annex=None, title_kind=None,
          organisation='', brand=None, cache=None, contents_heading=None, logo=None,
          review=False, bibliography=None, citation_style='apa', includes=(),
          columns=1):
    """Render one document to .docx. Returns build facts."""
    src = Path(source)
    work = Path(cache or src.parent) / ('.docx-%s' % src.stem)
    work.mkdir(parents=True, exist_ok=True)

    text = assemble.read(src, includes)
    front, text = front_mod.split(text)
    if review:
        front = front_mod.anonymise(front, prof)
    lines = text.split('\n')
    annex_lines = Path(annex).read_text(encoding='utf-8').split('\n') if annex else []
    # An embedded annex loses its own head - title, subtitle, its own contents -
    # in the reading edition, because it is folded into the parent document. The
    # first build here kept them and produced two headings the HTML did not
    # have, which is precisely what the cross-edition comparison is for.
    annex_title = None
    if annex_lines:
        # exactly what the reading edition does: the head runs to the first rule,
        # the annex keeps its title as a heading, and its subtitle and metadata
        # are dropped because it is folded into the parent document
        cut = next((i for i, l in enumerate(annex_lines)
                    if re.fullmatch(r'-{3,}', l.strip())), None)
        if cut is not None:
            annex_title = next((HEAD_RE.match(l.strip()).group(2)
                                for l in annex_lines[:cut] if l.strip().startswith('# ')), None)
            annex_lines = annex_lines[cut + 1:]

    start = None
    if contents_heading:
        marker = '## ' + contents_heading
        start = next((i for i, l in enumerate(lines) if l.strip().startswith(marker)), None)
    if start is None:
        start = next((i for i, l in enumerate(lines)
                      if re.fullmatch(r'-{3,}', l.strip())), 0) + 1
    head, body = lines[:start], lines[start:]
    kind, title, meta, _ = md.parse_head(head)

    images = {}
    if svgs:
        # the diagrams are already rasterised for the print edition; a Word file
        # cannot take the inline SVG the reading edition uses
        typst.rasterise(svgs, work)
        for i in range(len(svgs)):
            png = work / ('fig-%d.png' % i)
            if png.exists():
                images[i] = png

    doc = Document()
    if logo and Path(logo).exists():
        # Word cannot place an SVG, so one is rasterised the same way the
        # diagrams are - a project should not have to keep a second copy of its
        # own mark just because one of four editions is fussy
        mark = Path(logo)
        if mark.suffix.lower() == '.svg':
            typst.rasterise([mark.read_text(encoding='utf-8')], work, scale=4)
            raster = work / 'fig-0.png'
            mark = raster if raster.exists() else None
        if mark:
            doc.add_picture(str(mark), height=Mm(14))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    normal = doc.styles['Normal']
    # a project's own sans, if it names one: the first family in the stack, since
    # Word takes a single face rather than a fallback list
    # `-apple-system` and friends are CSS keywords, not typefaces: naming one in
    # a .docx gives Word a font it cannot find and a document that renders
    # differently on every machine
    SYSTEM_UI = {'-apple-system', 'blinkmacsystemfont', 'system-ui', 'ui-sans-serif',
                 'ui-serif', 'segoe ui'}
    stack = (brand or {}).get('sans') or (prof.get('fonts') or {}).get('sans') or ''
    families = [f.strip().strip('"\'') for f in stack.split(',')]
    normal.font.name = next((f for f in families
                             if f and f.lower() not in SYSTEM_UI), 'Calibri')
    normal.font.size = Pt(11)
    normal.font.color.rgb = _colour(brand, 'ink')
    if review:
        # the line numbers themselves go on at the end, once every section
        # exists - see _line_numbers()
        normal.paragraph_format.line_spacing = 2.0

    # not `text`: that name holds the document body, and shadowing it here made
    # the reference list search the title instead and silently find no citations
    for line, size, bold in ((kind or title_kind or '', 12, False),
                             (title or src.stem, 20, True)):
        if not line:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(line)
        run.bold, run.font.size = bold, Pt(size)
        # the kind above the title is amber on the cover of the reading and
        # print editions; here it was the body colour, so the one word that
        # says what the document is came out looking like the first line of it
        run.font.color.rgb = _navy(brand) if bold else _colour(brand, 'amber')
    if front.get('anonymised'):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(str(front['anonymised']))
        run.italic = True
        run.font.size = Pt(9)
    for name, marks in front_mod.byline(front):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(name)
        if marks:
            sup = para.add_run(marks)
            sup.font.superscript = True
    for key, value in sorted(front_mod.affiliations(front).items()):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mark = para.add_run(key)
        mark.font.superscript = True
        run = para.add_run(value)
        run.font.size = Pt(9)
    line = front_mod.corresponding(front, prof)
    if line:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(line).font.size = Pt(9)
    if front.get('abstract'):
        head = doc.add_paragraph()
        head.add_run(front_mod.label(prof, 'abstract')).bold = True
        _runs(doc.add_paragraph(), str(front['abstract']))
    if front.get('keywords'):
        para = doc.add_paragraph()
        para.add_run('%s: ' % front_mod.label(prof, 'keywords')).bold = True
        para.add_run(', '.join(str(k) for k in front['keywords']))
    for key, value in meta:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('%s: ' % key)
        run.bold = True
        run.font.color.rgb = _colour(brand, 'muted')
        _runs(para, value)
    if organisation:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(organisation).italic = True
    if columns > 1:
        # a section, not a page break: Word carries the column count on the
        # section, so the title block keeps the full measure and the body
        # opens in columns on the page after it
        _landscape(doc, False, columns)
    else:
        doc.add_page_break()

    figures = []
    label = prof['labels'].get('figure', 'Figure %d')
    part_banner = prof['structure'].get('part_banner')
    refs = xref.resolve(prof, body, annex_lines)
    convert(doc, body, figures, label, images, brand, part_banner, table=refs,
            columns=columns)
    if annex_lines:
        doc.add_page_break()
        if annex_title:
            heading = doc.add_heading(level=2)
            _runs(heading, annex_title)
            for run in heading.runs:
                run.font.color.rgb = _navy(brand)
        convert(doc, annex_lines, figures, prof['labels'].get('annex_figure', label),
                images, brand, part_banner, force_parts=True, table=refs,
                columns=columns)

    # The reference list. Without this the Word edition silently dropped it -
    # a submission copy with no bibliography - and only the cross-edition check
    # noticed, because the reading edition had one and Word did not.
    if bibliography:
        from . import citations as cite_mod
        keys = cite_mod.find(text + '\n' + '\n'.join(annex_lines))
        if keys:
            _, biblio = cite_mod.render(keys, bibliography, citation_style,
                                        prof['labels'].get('references', 'References'),
                                        prof.get('lang', 'en'))
            head = doc.add_heading(level=2)
            _runs(head, prof['labels'].get('references', 'References'))
            for item in re.findall(r'<li[^>]*>(.*?)</li>', biblio, re.S):
                plain = re.sub(r'<[^>]+>', '', item)
                plain = re.sub(r'\s+', ' ', ihtml.unescape(plain)).strip()
                if plain:
                    para = doc.add_paragraph(plain)
                    para.paragraph_format.left_indent = Mm(8)
                    para.paragraph_format.first_line_indent = Mm(-8)

    entries = front_mod.declarations(front, prof)
    if entries:
        head = doc.add_heading(level=2)
        _runs(head, front_mod.label(prof, 'declarations'))
        for key, value in entries:
            para = doc.add_paragraph()
            para.add_run('%s. ' % key).bold = True
            _runs(para, str(value))

    if review:
        _line_numbers(doc)
    doc.save(str(output))
    return {'figures': len(figures), 'tables': len(doc.tables),
            'headings': sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading')),
            'bytes': Path(output).stat().st_size}


def structure(path):
    """What a built .docx carries, for comparison with the reading edition."""
    doc = Document(str(path))
    heads = [p.text.strip() for p in doc.paragraphs
             if p.style.name.startswith('Heading') and p.text.strip()]
    images = sum(1 for r in doc.part.rels.values() if 'image' in r.reltype)
    return {'headings': heads, 'tables': len(doc.tables), 'figures': images}
