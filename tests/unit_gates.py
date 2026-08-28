#!/usr/bin/env python3
"""The gates' own rules, exercised directly.

A fixture proves the pipeline runs; it does not prove a rule refuses what it
claims to refuse. Everything here is documented behaviour - an unknown lint
pack is an error, a project profile layers over a shipped one, a deck warns
about a table nobody can read from the back of a room - reachable only by
feeding it the input a fixture never contains.
"""
import re
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from paperforge import (browser, citations, deck, docx, figures as fig_mod, lint,
                        matching,
                        markdown, pages,
                        profile, typst, verify, xref)

failures = []


def check(label, condition):
    print('  %-58s %s' % (label, 'ok' if condition else 'FAIL'))
    if not condition:
        failures.append(label)


def raises(label, fn, kind=Exception):
    try:
        fn()
    except kind as e:
        print('  %-58s ok (%s)' % (label, str(e)[:26]))
        return
    print('  %-58s FAIL (no error)' % label)
    failures.append(label)


def rules_fired(text, **kw):
    return {f['rule'] for f in lint.check_text(text, **kw)}


def main():
    print('lint rules')
    core = lint.ruleset()
    check('the core applies with no packs enabled', len(core) == len(lint.CORE))
    check('a pack adds its rules',
          len(lint.ruleset(['realtimex-loops'])) == len(lint.CORE) + len(lint.PACKS['realtimex-loops']))
    raises('an unknown pack is an error, not silently ignored',
           lambda: lint.ruleset(['no-such-pack']), ValueError)
    project = lint.ruleset([], [{'id': 'codename', 'pattern': 'PROJECT BLUEBIRD',
                                 'why': 'internal codename'}])
    check('a project rule is added', any(r[0] == 'codename' for r in project))
    check('a project rule blocks by default',
          next(r[1] for r in project if r[0] == 'codename') == 'block')

    check('an unfinished marker is caught', 'todo' in rules_fired('This is a TODO item'))
    check('placeholder text is caught', 'lorem' in rules_fired('Lorem ipsum dolor'))
    check('a filename shown to a reader is caught',
          'filename-label' in rules_fired('See [`REPORT.md`](./REPORT.md)'))
    check('an unrendered footnote is caught',
          'unsupported-footnote' in rules_fired('A claim.[^1]'))
    check('a caption is a supported construct now, not a blocked one',
          'unsupported-caption' not in rules_fired(': A caption {#fig-1}'))
    check('a project rule fires',
          'codename' in rules_fired('Regarding PROJECT BLUEBIRD.', rules=project))

    fenced = 'Text.\n```\nTODO fix this sample\n```\n'
    check('a rule cannot fire inside a code sample', 'todo' not in rules_fired(fenced))
    check('the same text outside a fence does fire',
          'todo' in rules_fired(fenced, skip_code=False))
    check('a finding reports its line',
          lint.check_text('ok\nok\nTODO here')[0]['line'] == 3)

    print('cross-references')
    with tempfile.TemporaryDirectory() as tmp:
        doc = Path(tmp) / 'r.md'
        doc.write_text('Drawn in @fig-flow.\n\n```mermaid\ngraph LR\nA-->B\n```\n\n'
                       ': How it flows {#fig-flow}\n\nAnd @fig-absent is nothing.\n',
                       encoding='utf-8')
        found = lint.check_references({'source_path': doc, 'annex_path': None},
                                      profile.load('en'))
        check('a reference to a label that does not exist is blocked',
              any(f['rule'] == 'dangling-reference' and f['match'] == '@fig-absent'
                  for f in found))
        check('a reference that resolves is not reported',
              not any(f['match'] == '@fig-flow' for f in found))
        check('the finding carries the line it is on',
              all(f['line'] > 0 for f in found if f['rule'] == 'dangling-reference'))

        twice = Path(tmp) / 'twice.md'
        twice.write_text(': One {#fig-a}\n\n: Two {#fig-a}\n', encoding='utf-8')
        found = lint.check_references({'source_path': twice, 'annex_path': None},
                                      profile.load('en'))
        check('a label declared twice is blocked, because every reference '
              'would mean the first',
              any(f['rule'] == 'duplicate-label' for f in found))

        # a document assembled from several files: the references have to
        # resolve across all of them, not only within the file they sit in
        one = Path(tmp) / 'one.md'
        two = Path(tmp) / 'two.md'
        one.write_text('```mermaid\ngraph LR\nA-->B\n```\n\n'
                       ': A chart {#fig-c}\n', encoding='utf-8')
        two.write_text('Discussed at @fig-c.\n', encoding='utf-8')
        found = lint.check_references(
            {'source_path': one, 'include_paths': [two], 'annex_path': None},
            profile.load('en'))
        check('a reference in one file to a label in another resolves',
              found == [])

    print('section labels')
    en = profile.load('en')
    lines = ['## Methods {#sec-methods}', '', 'Set out in @sec-methods.', '',
             '### Sub {.no-part #sec-sub}', '', 'Also @sec-sub and @sec-nowhere.', '',
             '```mermaid', 'graph LR', 'A-->B', '```', ': A chart {#fig-c}']
    table = xref.resolve(en, lines)
    check('a labelled heading enters the table as a section',
          table.get('sec-methods', {}).get('kind') == 'sec')
    # nothing in this pipeline numbers headings, and four emitters agreeing on
    # a counter is the failure xref.py exists to prevent
    check("a section has no number, and reads as the heading's own words",
          table.get('sec-methods', {}).get('number', 0) is None
          and table.get('sec-methods', {}).get('label') == 'Methods')
    check('the id is found whichever order the attribute is written in',
          table.get('sec-sub', {}).get('label') == 'Sub')
    check('a numbered kind alongside it still numbers',
          table.get('fig-c', {}).get('label') == 'Figure 1')
    check('a reference to a section resolves to the heading in prose',
          xref.substitute('Set out in @sec-methods.', table)
          == 'Set out in Methods.')
    check('a reference to a section that does not exist is dangling',
          [i for _, i in xref.dangling(lines, table)] == ['sec-nowhere'])
    check('a section id declared twice is reported',
          xref.duplicates(lines + ['## Again {#sec-methods}']) == ['sec-methods'])
    check('caption_of does not say "Methods. Methods"',
          'sec-methods' in table
          and xref.caption_of(table['sec-methods']) == 'Methods')
    # KINDS was declared here and consumed by nothing. The regexes derive from
    # it now, and this is the drift guard: it fires if a kind is added to one
    # and not the other. It passes on the old code too, where they agreed by
    # accident rather than by construction.
    check('every kind in KINDS is accepted by the reference syntax, and no other',
          all(xref.REF_RE.match('@%s-x' % k) for k in xref.KINDS)
          and not xref.REF_RE.match('@zzz-x'))
    check('and a section is not one of the numbered kinds',
          'sec' not in getattr(xref, 'NUMBERED', ('sec',))
          and 'sec' in xref.KINDS)

    # The attribute must not reach a reader in any edition. `take_equation`
    # records what happens when one does: `{#eq-x}` printed on the page.
    src = ('# Doc\n\n## Methods {#sec-methods}\n\nText.\n\n'
           '## Results {.no-part #sec-results}\n\nSet out in @sec-methods, and @sec-results.\n')
    resolved = 'Set out in Methods, and Results.'
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        (tmp / 'r.md').write_text(src, encoding='utf-8')
        markdown.build(tmp / 'r.md', tmp / 'r.html', prof=en)
        html = (tmp / 'r.html').read_text(encoding='utf-8')
        check('the reading edition keeps the id as the heading anchor',
              'id="sec-methods"' in html)
        check('and strips the attribute, and resolves the reference',
              '{#sec-methods}' not in html and '{.no-part' not in html
              and resolved in html)

        (tmp / 'd.md').write_text('# D\n\n---\n\n## Methods {#sec-methods}\n\n'
                                  'See @sec-methods.\n', encoding='utf-8')
        deck.build(tmp / 'd.md', tmp / 'd.html', prof=en)
        slides = (tmp / 'd.html').read_text(encoding='utf-8')
        check('the deck strips the attribute and resolves the reference',
              '{#sec-methods}' not in slides and 'See Methods.' in slides)

        typst.XREF.clear()
        typst.XREF.update(xref.resolve(en, src.split('\n')))
        typ = typst.convert(src.split('\n'), [], [], 'Figure %d', None)
        check('the print edition strips the attribute and resolves the reference',
              'sec-methods' not in typ and resolved in typ)

        docx.build(tmp / 'r.md', tmp / 'r.docx', en)
        import docx as _pd
        words = '\n'.join(p.text for p in _pd.Document(str(tmp / 'r.docx')).paragraphs)
        check('the Word edition strips the attribute and resolves the reference',
              '{#sec-methods}' not in words and '{.no-part' not in words
              and resolved in words)

    print('labels and headings with nothing pointing at them')
    # resolved defensively so a regression reads as failing checks rather than
    # a traceback that stops the rest of the suite
    orphans = getattr(lint, 'check_orphans', lambda *a, **k: [])
    stated = getattr(fig_mod, 'stated', lambda *a: set())
    with tempfile.TemporaryDirectory() as tmp:
        orph = Path(tmp) / 'r.md'
        orph.write_text(
            '## Part One {.part}\n\n### Chapter {#sec-ch}\n\n'
            'Prose that mentions @fig-used.\n\n'
            '```mermaid\ngraph LR\nA-->B\n```\n\n: Used chart {#fig-used}\n\n'
            '| a |\n| - |\n| 1 |\n\n: Lonely table {#tbl-lonely}\n\n'
            '## Empty\n\n## Last\n', encoding='utf-8')
        found = orphans({'source_path': orph, 'annex_path': None},
                        profile.load('en'))
        by = {(f['match'], f['rule']) for f in found}
        check('a float nothing refers to is reported',
              ('tbl-lonely', 'orphan-label') in by)
        check('one that is referred to is not', ('fig-used', 'orphan-label') not in by)
        # a heading exists to be a stable anchor as often as to be referred to,
        # so an unreferenced section label is not a finding
        check('an unreferenced section label is not reported',
              not any(m.startswith('sec-') for m, _ in by))
        check('a heading with nothing under it is reported',
              any(f['rule'] == 'empty-section' and 'Empty' in f['context'] for f in found))
        # the negative control that matters: every book in the corpus opens
        # parts with a banner heading whose content is the headings beneath it
        check('a part banner that opens deeper headings is not reported',
              not any(f['rule'] == 'empty-section' and 'Part One' in f['context']
                      for f in found))
        check('and everything reported is a warning, not a refusal',
              all(f['severity'] == 'warn' for f in found))

        # a reference from the annex still counts: the whole work is one document
        body = Path(tmp) / 'b.md'
        ann = Path(tmp) / 'a.md'
        body.write_text('| a |\n| - |\n| 1 |\n\n: Table {#tbl-x}\n', encoding='utf-8')
        ann.write_text('Discussed at @tbl-x.\n', encoding='utf-8')
        found = orphans({'source_path': body, 'annex_path': ann},
                        profile.load('en'))
        check('a float referred to only from the annex is not an orphan',
              not any(f['match'] == 'tbl-x' for f in found))

    print('figures declared and stated nowhere')
    with tempfile.TemporaryDirectory() as tmp:
        man = Path(tmp) / 'figures.toml'
        man.write_text('[[figure]]\nid = "target"\nlabel = "Target"\n'
                       'context = "target"\npattern = "\\\\d+%"\naccept = ["40%"]\n\n'
                       '[[figure]]\nid = "unused"\nlabel = "Unused"\n'
                       'context = "coverage"\npattern = "\\\\d+%"\naccept = ["7%"]\n',
                       encoding='utf-8')
        doc = Path(tmp) / 'd.md'
        doc.write_text('The target is 40% by 2030.\n', encoding='utf-8')
        declared = fig_mod.load(man)
        told = stated([doc], declared)
        check('a figure some document states is found stated', 'target' in told)
        check('one no document states is not', 'unused' not in told)
        # a disagreement is still a statement; reporting it as unused as well
        # would be two findings about one fact, the second of them false
        wrong = Path(tmp) / 'w.md'
        wrong.write_text('The target is 41% by 2030.\n', encoding='utf-8')
        check('a figure stated wrongly still counts as stated',
              'target' in stated([wrong], declared))

    print('claim labels')
    # resolved defensively so a regression reads as a failing check rather
    # than a traceback that stops the rest of the suite from running
    take_claim = getattr(xref, 'take_claim', lambda s: (s, None))
    claim_lines = ['## Methods {#sec-m}', '',
                   'The MLE is consistent under A1-A3. {#claim-mle}', '',
                   'A paragraph that ends in braces {like this}.', '',
                   'Stated again. {#claim-mle}']
    ctable = xref.resolve(en, claim_lines)
    check('a labelled paragraph enters the table as a claim',
          ctable.get('claim-mle', {}).get('kind') == 'claim')
    # a claim is the one labelled thing with no rendered form: not numbered,
    # not in the reference syntax, nothing on the page to point at
    check('a claim carries no number and no label to render',
          ctable.get('claim-mle', {}).get('number', 0) is None
          and ctable.get('claim-mle', {}).get('label') == '')
    check('a claim id declared twice is reported',
          xref.duplicates(claim_lines) == ['claim-mle'])
    check('a claim is not accepted by the reference syntax',
          not xref.REF_RE.match('@claim-mle')
          and 'claim' in getattr(xref, 'LABELLED', ())
          and 'claim' not in xref.KINDS)
    check('take_claim returns the paragraph without its label',
          take_claim('The MLE is consistent. {#claim-mle}')
          == ('The MLE is consistent.', 'claim-mle'))
    # a paragraph is entitled to end in braces, and stripping one that does
    # would delete an author's words
    check('a paragraph that merely ends in braces is left alone',
          take_claim('ends in braces {like this}.')
          == ('ends in braces {like this}.', None))

    check('a claim referred to in prose is blocked',
          'claim-reference' in rules_fired('As argued in @claim-mle, this holds.'))
    check('the definition itself is not',
          'claim-reference' not in rules_fired('The MLE is consistent. {#claim-mle}'))
    check('and a section reference is still fine',
          'claim-reference' not in rules_fired('As set out in @sec-methods.'))

    # The label must not reach a reader in any edition. Same trap as `{#eq-x}`,
    # which take_equation records printing on the page.
    claim_src = ('# Doc\n\n## Methods {#sec-m}\n\n'
                 'The MLE is consistent under A1-A3. {#claim-mle}\n\n'
                 'A paragraph that ends in braces {like this}.\n')
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        (tmp / 'c.md').write_text(claim_src, encoding='utf-8')
        markdown.build(tmp / 'c.md', tmp / 'c.html', prof=en)
        html = (tmp / 'c.html').read_text(encoding='utf-8')
        check('the reading edition strips the label and keeps the prose',
              'claim-mle' not in html and 'consistent under A1-A3.</p>' in html
              and 'braces {like this}.' in html)

        (tmp / 'cd.md').write_text('# D\n\n---\n\n## S\n\n'
                                   'The MLE is consistent. {#claim-mle}\n', encoding='utf-8')
        deck.build(tmp / 'cd.md', tmp / 'cd.html', prof=en)
        check('the deck strips the label',
              'claim-mle' not in (tmp / 'cd.html').read_text(encoding='utf-8'))

        typst.XREF.clear()
        typst.XREF.update(xref.resolve(en, claim_src.split('\n')))
        typ = typst.convert(claim_src.split('\n'), [], [], 'Figure %d', None)
        check('the print edition strips the label and keeps the braces',
              'claim-mle' not in typ and 'like this' in typ)

        docx.build(tmp / 'c.md', tmp / 'c.docx', en)
        import docx as _pd
        words = '\n'.join(q.text for q in _pd.Document(str(tmp / 'c.docx')).paragraphs)
        check('the Word edition strips the label and keeps the prose',
              'claim-mle' not in words
              and 'The MLE is consistent under A1-A3.' in words
              and 'braces {like this}.' in words)

    print('display equations')
    from paperforge import xref as xr
    block = ['$$', 'a^2 + b^2 = c^2', '$$ {#eq-p}', '', 'after']
    expr, ident, after = xr.take_equation(block, 0)
    check('a labelled block yields its expression and id',
          expr == 'a^2 + b^2 = c^2' and ident == 'eq-p')
    check('and the fence line is consumed, so the label cannot print',
          after == 3 and block[after] == '')
    check('an unlabelled block still yields its expression',
          xr.take_equation(['$$', 'x', '$$'], 0)[:2] == ('x', None))
    check('a line that is not a fence is left alone',
          xr.take_equation(['ordinary prose'], 0) == (None, None, 0))
    check('an unterminated fence is left to the paragraph path',
          xr.take_equation(['$$', 'x'], 0) == (None, None, 0))
    numbered = xr.resolve(profile.load('en'), block)
    check('the equation is numbered like any other labelled thing',
          numbered['eq-p']['label'] == 'Equation 1')

    print('whether a PDF can be read back, as opposed to how much of it there is')
    # This counted characters once, and a character count is a claim about
    # volume being read as a claim about legibility. Measured across the
    # fixtures, volume ran from 0.07 to 44 and separated nothing, while
    # correspondence ran 0.75-0.97 for documents whose print checks work and
    # 0.00-0.08 for those whose do not. The floor sits in that gap.
    prose = ('The estimator is consistent under assumptions A1 to A3. Concentration '
             'risk follows from refining capacity rather than extraction. Policy '
             'attention should follow the midstream instead of the mine.')
    good = pages.correspondence(prose, prose)
    check('text that reads back as itself is usable', good['usable'] and good['ratio'] == 1.0)

    # the Arabic case, verbatim: the whole document comes back, shaped into
    # presentation forms and in visual order, so all of it is there and none
    # of it matches. Volume said 103% and the old check waved it through.
    shaped = 'يﺬﯿﻔﻨﺘﻟا ﺺﺨﻠﻤﻟا ' * 40
    arabic = pages.correspondence('الملخص التنفيذي والسياق والنتائج والخاتمة ' * 8,
                                  shaped, fold_diacritics=False)
    check('text that is all there and matches nothing is refused',
          not arabic['usable'])
    check('and volume alone would have passed it', arabic['volume'] > 0.45)
    check('the refusal says what is wrong', 'does not match' in (arabic['why'] or ''))

    # untestable is never passed, and never quietly failed either
    short = pages.correspondence('a b c. 12 34', 'anything at all')
    check('a document with too few distinctive words declines, with a reason',
          not short['usable'] and 'too few' in (short['why'] or ''))

    print('a contents entry and the heading it names, in every script')
    # `markdown._norm` was a second copy of the normaliser, and it was the
    # ASCII-only one - `[^a-z0-9\\s.:]` erased Arabic and Chinese completely, so
    # both sides of the match reduced to their punctuation and no entry could
    # ever be numbered. `profile.normalise` had already been fixed for exactly
    # this and says so in its docstring; this copy never heard, which is the
    # same shape as the threshold in matching.py.
    #
    # Checked for every shipped profile rather than the two that broke, so a
    # fifth language cannot arrive with the same hole.
    named = {'ar': ('1. الملخص التنفيذي', 'الملخص التنفيذي'),
             'en': ('1. Executive summary', 'Executive summary'),
             'vi': ('1. Tóm tắt', 'Tóm tắt'),
             'zh': ('1. 执行摘要', '执行摘要')}
    shipped = sorted(q.stem for q in
                     (Path(__file__).resolve().parents[1] / 'paperforge/profiles').glob('*.toml'))
    check('every shipped profile has a case here', sorted(named) == shipped)
    for lang in shipped:
        entry, heading = named[lang]
        markdown.PROF = profile.load(lang)
        shared = markdown.entry_keys(entry) & markdown.entry_keys(heading)
        check('%s: the contents entry and its heading share a key' % lang, bool(shared))
    # not merely unmatched but indistinguishable: the ASCII-only signature
    # collapsed every Arabic heading to the same empty key
    markdown.PROF = profile.load('ar')
    check('and two different headings do not collapse to one key',
          markdown.entry_keys('الملخص التنفيذي') != markdown.entry_keys('السياق'))

    print('everything this pipeline says it supports, built by something')
    # Three defects this session were the same shape: a capability shipped, was
    # documented as supported, and was exercised by nothing. `deck.build` was
    # reachable only by hand and had three defects; the `ar` profile had never
    # been built at all and a scaffolded Arabic project failed `all`. The
    # doctrine section calls this out for emitters - it is broader than that.
    # A profile is a capability. A document type is a capability.
    #
    # Anything deliberately uncovered belongs here with its reason, so a second
    # one has to be argued for in the diff that adds it rather than noticed
    # later by somebody scaffolding a project.
    UNCOVERED = {}

    repo = Path(__file__).resolve().parents[1]
    from paperforge import cli as cli_mod
    shipped = {'profile': {q.stem for q in (repo / 'paperforge/profiles').glob('*.toml')},
               'type': set(cli_mod.BUILTIN_TYPES)}
    used = {'profile': set(), 'type': set()}
    for manifest in sorted(repo.glob('tests/**/documents.toml')):
        text = manifest.read_text(encoding='utf-8')
        used['profile'] |= set(re.findall(r'profile\s*=\s*"([\w-]+)"', text))
        used['type'] |= set(re.findall(r'type\s*=\s*"([\w-]+)"', text))
        if re.search(r'^\s*\[\[collection\.document\]\]', text, re.M):
            used['type'].add('report')          # the default when none is named

    for kind in ('profile', 'type'):
        missing = sorted(shipped[kind] - used[kind] - set(UNCOVERED))
        check('every shipped %s is built by some fixture' % kind, missing == [])
        if missing:
            print('          uncovered: %s' % ', '.join(missing))
    check('and nothing is excused without a reason written down',
          all(v.strip() for v in UNCOVERED.values()))

    print('a threshold may never exceed the pool it draws from')
    quorum = getattr(matching, 'quorum', lambda pool, floor: max(floor, pool - 1))
    over = [(pool, floor) for pool in range(0, 30) for floor in range(1, 6)
            if quorum(pool, floor) > pool]
    # the property, not the three instances: this is what all three had wrong
    check('no pool and floor produce a threshold nothing could meet', over == [])
    check('and the floor is still respected where there is room',
          quorum(9, 3) == 8 and quorum(2, 3) == 2)

    # the regression that made it worth extracting: a two-word contents entry
    # whose words are on the page but not adjacent, which is what pdfplumber
    # produces from a two-column page - see "a check that reads an artifact"
    def entry_verdict(label, text):
        words = [w for w in label.split() if len(w) > 1 and not w.isdigit()][:5]
        if len(' '.join(words)) < 10:
            return 'skip'
        if ' '.join(words) in text or sum(w in text for w in words) >= quorum(len(words), 3):
            return 'confirmed'
        return 'wrong'
    check('a two-word entry present but split is confirmed, not reported wrong',
          entry_verdict('Global demand', 'Global 12 demand rose') == 'confirmed')
    check('and one that is genuinely half-absent is still wrong',
          entry_verdict('Global demand', 'Global rose') == 'wrong')

    # a source scan, so a fourth copy of the arithmetic cannot arrive quietly
    unclamped = []
    for path in sorted(Path(__file__).resolve().parents[1].glob('paperforge/*.py')):
        if path.name == 'matching.py':
            continue
        for n, line in enumerate(path.read_text(encoding='utf-8').split('\n'), 1):
            if re.search(r'max\(\s*\d+\s*,\s*len\(', line) and not line.lstrip().startswith('#'):
                unclamped.append('%s:%d' % (path.name, n))
    check('no module computes a bare max(n, len(...)) threshold of its own',
          unclamped == [])

    print('what a finding is allowed to say')
    # resolved defensively so a regression reads as failing checks, not a
    # traceback that stops the rest of the suite
    vocabulary = getattr(lint, 'SEVERITIES', ())
    check('the vocabulary is exactly four words',
          vocabulary == ('block', 'manual', 'warn', 'skip'))
    graded = [{'severity': s, 'rule': s}
              for s in vocabulary or ('block', 'warn')]
    s = lint.summarise(graded)
    check('every severity is counted',
          s.get('counts') == {'block': 1, 'manual': 1, 'warn': 1, 'skip': 1})
    # manual and skip are not softer blocks: neither stops publication, and the
    # whole point of naming them is that they are not warnings either
    check('and only a block blocks', s['blocking'] == 1)
    raises('a severity outside the set is refused rather than counted as a warning',
           lambda: lint.summarise([{'severity': 'urgent', 'rule': 'x'}]), ValueError)

    # a count nobody can chase is not a report: an untestable entry carries the
    # reason it could not be tested, which is what makes it a skip and not a pass
    empty = pages.audit(__file__, __file__, 'nothing-here', '', '', '')
    check('an entry that cannot be tested is a list of reasons, not a tally',
          empty['untestable'] == [] and isinstance(empty['untestable'], list))

    print('the publication allowlist')
    declared, blocked, embedded = {'report.md'}, {'REVIEW.md'}, {'annex.md'}
    check('a declared document passes',
          lint.check_publishable('/x/report.md', declared, blocked, embedded) == [])
    check('an embedded annex is neither publishable nor an error',
          lint.check_publishable('/x/annex.md', declared, blocked, embedded) == [])
    check('an internal file is refused',
          lint.check_publishable('/x/REVIEW.md', declared, blocked, embedded)[0]['rule']
          == 'not-publishable')
    check('an undeclared file is refused',
          lint.check_publishable('/x/stray.md', declared, blocked, embedded)[0]['rule']
          == 'undeclared')

    print('profiles')
    check('every shipped profile loads',
          all(profile.load(n).get('lang') for n in profile.available()))
    check('Vietnamese folds to a matchable form',
          profile.normalise('PHẦN III: BỐI CẢNH') == 'phan iii: boi canh')
    check('folding can be refused, because some marks are letters',
          profile.normalise('PHẦN', fold_diacritics=False) != 'phan')
    merged = profile.merge({'labels': {'figure': 'Figure %d', 'annex_badge': 'Annex'}},
                           {'labels': {'figure': 'Diagram %d'}})
    check('an override merges into a nested table, it does not replace it',
          merged['labels'] == {'figure': 'Diagram %d', 'annex_badge': 'Annex'})
    with tempfile.TemporaryDirectory() as tmp:
        local = Path(tmp) / 'profile.toml'
        local.write_text('[labels]\nfigure = "Rajah %d"\n', encoding='utf-8')
        layered = profile.load_file(local, profile.load('en'))
        check('a project profile layers over a shipped one',
              layered['labels']['figure'] == 'Rajah %d'
              and layered['structure']['contents_heading'] == 'CONTENTS')
        raises('a missing profile file is an error, not a default',
               lambda: profile.load_file(Path(tmp) / 'absent.toml', None))

    print('decks')
    source = ['## First slide', 'Body.', '', '---', '', '## Second slide',
              '> notes: say this out loud', 'More body.']
    cut = deck.slides(source)
    check('a ## heading and an explicit rule both start a slide', len(cut) == 2)
    check('a slide keeps its own body', 'Body.' in cut[0])
    body, notes = deck._notes(cut[1])
    check('speaker notes are pulled out of the slide body',
          notes == ['say this out loud'])
    check('the note is not left in the visible text',
          not any('notes:' in l for l in body))
    check('words are counted by spaces for Latin', deck.count_words('one two three') == 3)
    # ~2.2 characters carry a word, so a dense Chinese slide does not register
    # as three "words" and slip past the length check
    check('CJK is counted by character, having no word spaces',
          deck.count_words('关键矿产供应链', units='characters') == 3)
    check('the same string counted by spaces looks like one word',
          deck.count_words('关键矿产供应链') == 1)

    dense = '<section><table>%s</table></section>' % ('<tr><td>x</td></tr>' * 9)
    check('a table nobody could read from the back of a room is reported',
          any('table' in w for w in deck.audit(dense)))
    wordy = '<section><p>%s</p></section>' % ('word ' * 140)
    check('an overfull slide is reported', any('word' in w for w in deck.audit(wordy)))
    check('a reasonable slide is not reported',
          deck.audit('<section><p>Three short words</p></section>') == [])

    # A deck renders through markdown.py, whose XREF is a module global that
    # only markdown.build ever filled. So a deck built alone printed `@tbl-t`
    # to the reader and silently dropped its own captions, and a deck built
    # after a report resolved against *the report's* table - a plausible
    # "Table 1" on a slide with no tables, depending on manifest order.
    with tempfile.TemporaryDirectory() as tmp:
        tmp, en = Path(tmp), profile.load('en')
        alone = tmp / 'alone.md'
        alone.write_text('# Deck\n\n---\n\n## One\n\nSee @tbl-t.\n\n'
                         '| a | b |\n| - | - |\n| 1 | 2 |\n\n: A table {#tbl-t}\n',
                         encoding='utf-8')
        deck.build(alone, tmp / 'alone.html', prof=en)
        html = (tmp / 'alone.html').read_text(encoding='utf-8')
        check('a deck resolves a label declared in its own source', 'Table 1' in html)
        check('so no reference reaches the slide as raw source', '@tbl-t' not in html)
        check("and the deck's own caption is rendered rather than dropped",
              'A table' in html)

        report = tmp / 'report.md'
        report.write_text('# Report\n\n## S\n\n| a |\n| - |\n| 1 |\n\n'
                          ': Report table {#tbl-r}\n', encoding='utf-8')
        markdown.build(report, tmp / 'report.html', prof=en)
        after = tmp / 'after.md'
        after.write_text('# Deck\n\n---\n\n## One\n\nSee @tbl-r.\n', encoding='utf-8')
        deck.build(after, tmp / 'after.html', prof=en)
        check('a deck does not inherit the table of the document built before it',
              'Table 1' not in (tmp / 'after.html').read_text(encoding='utf-8'))

    print('citations')
    check('a single key is found', citations.find('A claim [@nq57].') == ['nq57'])
    check('several keys in one marker are found',
          citations.find('Both [@a; @b] agree.') == ['a', 'b'])
    check('a repeated key is reported once',
          citations.find('[@a] and again [@a]') == ['a'])
    check('an email address is not a citation', citations.find('write to a@b.com') == [])
    with tempfile.TemporaryDirectory() as tmp:
        bib = Path(tmp) / 'refs.bib'
        bib.write_text('@legislation{nq57,\n  title = {A decree},\n  year = {2026},\n}\n'
                       '@report{ok,\n  title = {A report},\n  year = {2026},\n}\n',
                       encoding='utf-8')
        warned = dict(citations.dangling_dates(bib))
        check('a year-only @legislation entry is flagged for the stray comma',
              warned.get('nq57') == 'legislation')
        check('@report with the same year is not flagged', 'ok' not in warned)
        check('the warning can be limited to the keys actually cited',
              citations.dangling_dates(bib, {'ok'}) == [])

    print('coverage of lines that render as separate spans')
    # Metadata renders across styled spans, so the contiguous probe legitimately
    # fails and a word-presence fallback decides. Its floor used to be a flat
    # max(2, len(words) - 1), which can ask for more matches than there are
    # words: words of one or two characters are dropped, so `**Prepared by:**
    # Le` reduces to ["Prepared"] and could never reach two. "RTA" passes and
    # "Le" does not - a fair way to lose a Vietnamese name - and no document
    # could satisfy the gate either way.
    with tempfile.TemporaryDirectory() as tmp:
        cov = Path(tmp) / 'r.md'
        cov.write_text('**Prepared by:** Le\n\n**Reviewed by:** Vu\n\n'
                       'A whole sentence of ordinary prose that should be found.\n\n'
                       'Another sentence of ordinary prose that is absent.\n',
                       encoding='utf-8')
        # the real markup: adjacent spans, no whitespace and no colon between
        # them, which is exactly why the contiguous probe cannot match and the
        # word fallback is reached at all
        rendered = ('<html><body><main>'
                    '<div class="meta-item"><span class="meta-k">Prepared by</span>'
                    '<span class="meta-v">Le</span></div>'
                    '<p>A whole sentence of ordinary prose that should be found.</p>'
                    '</main></body></html>')
        gone = [m[2] for m in verify.coverage(rendered, cov)]
        check('a two-character metadata value that did render is not reported',
              not any('Prepared' in g for g in gone))
        # the fix must not buy that by making the check vacuous
        check('one that did not render still is',
              any('Reviewed' in g for g in gone))
        check('and ordinary prose that did not render still is',
              any('absent' in g for g in gone))

    print('structural checks on a document built to be broken')
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / 'report.md'
        src.write_text('# T\n## Title\n\n---\n\n'
                       'A sentence that reaches the rendered page unchanged.\n\n'
                       'A second sentence that the renderer will lose entirely.\n',
                       encoding='utf-8')
        page = Path(tmp) / 'report.html'
        page.write_text(
            '<html><body><main>'
            '<h2 id="one">Title</h2>'
            '<p>A sentence that reaches the rendered page unchanged.</p>'
            '<a href="#nowhere">a reference to nothing</a>'
            '<img src="https://example.com/logo.png">'
            '<a href="https://ifr.org/worldrobotics">a source the reader may follow</a>'
            '<b><i>crossed</b></i>'
            '</main></body></html>', encoding='utf-8')
        r = verify.check(page, src)
        check('crossed tags are reported',
              any('expected </i>, got </b>' in e for e in r['markup_errors']))
        check('an anchor pointing nowhere is reported', r['broken_anchors'] == ['nowhere'])
        # A remote image is a dependency: without the network the page is
        # wrong. A remote link is a citation: the page is identical offline.
        # Blocking on both meant a reference list with retrieval URLs, which is
        # what a reference list is for, could not pass.
        check('a remote asset is reported as a dependency',
              r['external_assets'] == ['https://example.com/logo.png'])
        check('a link the reader may follow is not',
              r['external_links'] == ['https://ifr.org/worldrobotics'])
        check('a source line that never reached the page is reported',
              any('second sentence' in m[2] for m in r['missing_content']))
        check('a line that did reach the page is not reported',
              not any('reaches the rendered page' in m[2] for m in r['missing_content']))

        page.write_text('</p><html><body><main><p>text</p></main></body></html>',
                        encoding='utf-8')
        check('a closing tag with nothing open is reported as stray',
              any('stray' in e for e in verify.check(page, src)['markup_errors']))
        page.write_text('<html><body><main><div><p>text', encoding='utf-8')
        left = verify.check(page, src)['unclosed']
        check('tags left open at the end of the file are reported',
              {'div', 'p'} <= set(left))
        page.write_text('<html><body><main><p>text</p></main></body></html>', encoding='utf-8')
        clean = verify.check(page, src)
        check('a well-formed document reports nothing',
              not clean['markup_errors'] and not clean['unclosed'])

    print('raw markup that reached the page')
    check('an unrendered tag is a leak',
          any(l['match'] == '<br>' for l in verify.leaks('a TRỊ<br>(Verified)')))
    check('a real HTML entity is a leak',
          any(l['match'] == '&nbsp;' for l in verify.leaks('a&nbsp;b')))
    check('a numeric entity is a leak', any(l['kind'] == 'entity' for l in verify.leaks('a&#160;b')))
    check('"KH&CN;" is prose, not an entity', verify.leaks('KH&CN; and R&D; costs') == [])
    check('"<10%" and ">55" are prose, not tags', verify.leaks('Nhóm 2: <10% and >55–60%') == [])
    check('unrendered emphasis is a leak', any(l['kind'] == 'emphasis' for l in verify.leaks('a **b**')))

    print('branding')
    from paperforge import markdown as md
    prof = profile.load('en')
    plain = md.theme_override(prof, None)
    check('an unbranded project is given no palette at all',
          '--navy' not in plain and '--navy-deep' not in plain)
    # only what differs from the shipped defaults: the English profile's serif
    # *is* the default, so re-declaring it would be a copy of the line above it
    check("and only the faces that differ from the stylesheet's",
          '--sans' in plain and '--serif' not in plain)
    check('a profile whose faces both differ declares both',
          all(tok in md.theme_override(profile.load('zh'), None)
              for tok in ('--sans', '--serif')))
    branded = md.theme_override(prof, {'navy': '#5b2333', 'bg': '#f7f4ef'})
    check('a declared palette is emitted as tokens',
          '--navy: #5b2333' in branded and '--bg: #f7f4ef' in branded)
    check('together with the shades derived from it, which were not declared',
          '--navy-deep:' in branded and '--navy-tint:' in branded)
    housed = md.theme_override(prof, {'serif': 'Palatino, serif'})
    check('a project may name its own face, overriding the profile',
          '--serif: Palatino, serif' in housed)
    check('and the profile still supplies the one it did not name',
          '--sans' in housed)

    with tempfile.TemporaryDirectory() as tmp:
        mark = Path(tmp) / 'mark.svg'
        mark.write_text('<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 10 10">'
                        '<rect width="10" height="10"/></svg>', encoding='utf-8')
        tag = md.logo_tag(mark)
        check('an SVG mark is inlined as markup, not linked',
              '<svg' in tag and 'src=' not in tag)
        raster = Path(tmp) / 'mark.png'
        raster.write_bytes(b'\x89PNG\r\n\x1a\n' + b'\0' * 40)
        check('any other format is inlined as a data URI',
              'data:image/png;base64,' in md.logo_tag(raster))
        check('a mark that is not there is not a broken image',
              md.logo_tag(Path(tmp) / 'absent.svg') == '')

    print('locating the contents in a printed PDF')
    # These two guards were added after a real defect: with every contents entry
    # shorter than the length filter, the search scored every page zero and
    # returned the cover, exempting the wrong page from the near-empty check.
    toc = '<h2 id="contents">CONTENTS</h2><ol><li>1. Context</li>' \
          '<li>2. Sources and method</li><li>3. Conclusions drawn</li></ol><h2 id="x">X</h2>'
    body = ['cover page only', 'contents 1. context 2. sources and method 3. conclusions drawn',
            'context the body of the report begins here']
    check('the contents page is found, not the cover',
          pages.contents_pages([pages.norm(t) for t in body], toc, 'contents') == {1})
    short = '<h2 id="contents">CONTENTS</h2><ol><li>1. A</li><li>2. B</li></ol><h2 id="x">X</h2>'
    found = pages.contents_pages([pages.norm(t) for t in
                                  ['cover', 'contents 1. a 2. b', 'a the body']], short, 'contents')
    check('entries too short for the length filter still locate the contents', found == {1})
    empty = '<h2 id="contents">CONTENTS</h2><ol></ol><h2 id="x">X</h2>'
    check('a contents with no entries exempts nothing rather than guessing',
          pages.contents_pages([pages.norm(t) for t in body], empty, 'contents') == set())
    check('a contents that cannot be located exempts nothing',
          pages.contents_pages([pages.norm('unrelated text'), pages.norm('more prose')],
                               toc, 'contents') == set())

    print('finding a browser')
    real_which = browser.shutil.which
    try:
        browser.shutil.which = lambda name: '/usr/bin/%s' % name if name == 'chromium' else None
        check('a browser on PATH is used', browser.chrome() == '/usr/bin/chromium')
        browser.shutil.which = lambda name: None
        real_candidates = browser.CANDIDATES
        browser.CANDIDATES = []
        try:
            browser.chrome()
            check('no browser anywhere is an error, not a silent blank render', False)
        except RuntimeError:
            check('no browser anywhere is an error, not a silent blank render', True)
        finally:
            browser.CANDIDATES = real_candidates
    finally:
        browser.shutil.which = real_which

    # The document <title> is the browser tab, the bookmark and the share
    # preview - the one string a reader meets before the document. It carried
    # "Báo cáo nghiên cứu chính sách" for every language this pipeline ships,
    # because it was a literal in the template and no profile could reach it.
    # The rule is the general one: nothing inside <title> may be text.
    # Typst markup that must survive as text. A character missing from this
    # list is not a rendering error - it is a silent deletion, and the reading
    # edition keeps what the print edition drops.
    print('typst escaping')
    from paperforge import typst as typst_mod
    for ch in '#$*_`<>@~\\':
        check('%r is escaped, not read as markup' % ch,
              typst_mod.esc(ch) == '\\' + ch)
    check('a tilde survives a whole cell',
          typst_mod.esc('~28x') == '\\~28x')

    print('the document head')
    import re as _re
    template = (Path(__file__).resolve().parents[1]
                / 'paperforge/theme/document.html').read_text(encoding='utf-8')
    m = _re.search(r'<title>(.*?)</title>', template, _re.S)
    check('the template has a title', bool(m))
    literal = _re.sub(r'\{\{\w+\}\}', '', m.group(1) if m else '').strip(' \u2014\u2013-|\u00b7')
    check('and no literal text in it, in any language', literal == '')

    if failures:
        print('\n%d check(s) failed: %s' % (len(failures), '; '.join(failures)))
        return 1
    print('\ngates: all checks passed')
    return 0


if __name__ == '__main__':
    sys.exit(main())
