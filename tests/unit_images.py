#!/usr/bin/env python3
"""An image as a figure: found, numbered, inlined, or refused.

This whole path existed as documentation and nothing else. `![alt](src)` was
matched by the link pattern, which took the bracket half and left the bang
behind as text; the caption under it was consumed by nothing and printed to the
reader with its `{#fig-x}` braces showing; and `@fig-x` still resolved to a
number for a float that was never on the page. Every gate passed.

So the tests here are mostly about what must now *fail*, and about the one
number that has to stay right: a figure's ordinal counts floats, while the
diagram raster index counts diagrams, and an image between two diagrams drew
the wrong one while those were the same counter.
"""
import base64
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from docx import Document

from paperforge import docx as docx_mod
from paperforge import images, lint, markdown, profile, require, typst, verify, xref

failures = []

# a real 8x8 PNG, not a plausible-looking one. The first constant here decoded
# as bytes and inlined as a data URI without complaint, and only Typst - which
# actually decodes what it places - reported the CRC error.
PNG = base64.b64decode(
    'iVBORw0KGgoAAAANSUhEUgAAAAgAAAAICAIAAABLbSncAAAAEUlEQVR4nGPwSG3EihiGlgQA'
    'cjtLgZBs74EAAAAASUVORK5CYII=')
SVG = '<svg viewBox="0 0 10 10" style="max-width:10px"><g/></svg>'

FRONT_DOC = '''+++
abstract = "A summary with ![a mark](f.png) inside it."
+++

# DOCUMENT
## Title

---
**Prepared by:** Test

---

## Body

Some prose long enough to be a paragraph in a document.
'''


def check(label, condition):
    print('  %-58s %s' % (label, 'ok' if condition else 'FAIL'))
    if not condition:
        failures.append(label)


def rules(found):
    return {(f['rule'], f['severity']) for f in found}


def render(lines, root, svgs=(), prof=None):
    """Drive the HTML emitter the way build() does, without a whole project."""
    prof = prof or profile.load('en')
    markdown.PROF = prof
    markdown.SVGS[:] = list(svgs)
    markdown.SRC['dir'] = Path(root)
    markdown.XREF.clear()
    markdown.XREF.update(xref.resolve(prof, lines))
    markdown.FIG.update(n=0, base=0, dgm=0, label=prof['labels']['figure'])
    return markdown.convert(lines, [])


def document(root, body, name='doc.md'):
    path = Path(root) / name
    path.write_text('\n'.join(body), encoding='utf-8')
    return {'source_path': str(path), 'include_paths': (), 'annex_path': None}


def main():
    print('reading image references out of a source')
    lines = ['text ![one](a.png) more', '', '![two](b/c.svg)', '',
             '```', '![not this](d.png)', '```']
    found = images.refs(lines)
    check('both real references are found, in order',
          [(n, s) for n, _, s in found] == [(1, 'a.png'), (3, 'b/c.svg')])
    check('a reference inside a fence is code, not an image', len(found) == 2)

    print('resolving a path')
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'figures').mkdir()
        (root / 'figures' / 'f.png').write_bytes(PNG)
        check('relative to the document, not the working directory',
              images.resolve('figures/f.png', root) == root / 'figures' / 'f.png')
        check('a path with no file resolves to nothing',
              images.resolve('figures/gone.png', root) is None)
        check('a remote src resolves to nothing rather than being fetched',
              images.resolve('https://example.com/f.png', root) is None)
        check('a protocol-relative src is remote too',
              images.resolve('//example.com/f.png', root) is None)
        uri = images.data_uri(root / 'figures' / 'f.png')
        check('the file is carried in the document, not linked',
              uri.startswith('data:image/png;base64,')
              and base64.b64decode(uri.split(',', 1)[1]) == PNG)

    print('an image on its own line is a figure')
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'f.png').write_bytes(PNG)
        body = ['## Findings {#sec-f}', '', 'See @fig-photo.', '',
                '![A field](f.png)', '', ': A salinity-affected field. {#fig-photo}']
        html = render(body, root)
        check('it renders as a figure, anchored on its label',
              '<figure class="diagram plate" id="fig-photo">' in html)
        check('the picture is inlined, not linked',
              'src="data:image/png;base64,' in html)
        check('the caption is consumed and numbered',
              '<figcaption>Figure 1. A salinity-affected field.' in html)
        check('the caption does not also print as prose',
              '{#fig-photo}' not in html and '<p>: A' not in html)
        check('no literal bang survives into the page', '!<a href' not in html)
        check('the reference resolves to the figure', 'See Figure 1.' in html)

        print('an image inside a sentence stays inside the sentence')
        html = render(['Some prose ![a mark](f.png) mid-line.'], root)
        check('it renders as an image', '<img src="data:image/png' in html)
        check('and not as a bang and a link', '!<a href' not in html)
        check('and does not become a figure', 'figcaption' not in html)

        print('syntax shown to a reader is not a picture placed for one')
        html = render(['Write `![alt](f.png)` to place one.'], root)
        check('an image in a code span stays code',
              '<code>![alt](f.png)</code>' in html and '<img' not in html)
        check('and lint does not go looking for the file',
              images.refs(['Write `![alt](f.png)` to place one.']) == [])

        print('a float ordinal counts floats; a raster index counts diagrams')
        body = ['![A field](f.png)', '', ': The photograph. {#fig-photo}', '',
                '```mermaid', 'graph TD', 'A-->B', '```', '',
                ': The diagram. {#fig-flow}']
        html = render(body, root, svgs=[SVG])
        check('the image is Figure 1', '<figcaption>Figure 1. The photograph.' in html)
        check('the diagram after it is Figure 2',
              '<figcaption>Figure 2. The diagram.' in html)
        check('the diagram still draws the first rendered SVG', SVG in html)
        check('the diagram count is diagrams, not floats', markdown.FIG['dgm'] == 1)

        print('a number counts floats, not captions')
        # the collision this rule exists for: the emitters number every float
        # positionally and print "Figure 1" under an uncaptioned one, while the
        # label table used to number captions 1..N. One of each gave two of them
        body = ['```mermaid', 'graph TD', 'A-->B', '```', '',
                '```mermaid', 'graph TD', 'C-->D', '```', '',
                ': The second diagram. {#fig-second}', '', 'See @fig-second.']
        table = xref.resolve(profile.load('en'), body)
        check('a captioned figure is numbered past the uncaptioned one before it',
              table['fig-second']['number'] == 2)
        html = render(body, root, svgs=[SVG, SVG])
        check('and the page agrees with the table',
              '<figcaption>Figure 1<' in html
              and '<figcaption>Figure 2. The second diagram.' in html)
        check('so the reference points at the right one', 'See Figure 2.' in html)

        body = ['![A field](f.png)', '', '![Another](f.png)', '',
                ': The second image. {#fig-b}']
        check('an uncaptioned image counts the same way',
              xref.resolve(profile.load('en'), body)['fig-b']['number'] == 2)

        body = ['| a | b |', '|---|---|', '| 1 | 2 |', '', '| c | d |',
                '|---|---|', '| 3 | 4 |', '', ': The second table. {#tbl-b}']
        check('an uncaptioned table takes no number: nothing prints one for it',
              xref.resolve(profile.load('en'), body)['tbl-b']['number'] == 1)

        # every place a float can hide from the scanner but not from an emitter
        body = ['> [!note]', '> ![a](f.png)', '', '```mermaid', 'graph TD',
                'A-->B', '```', '', ': The diagram. {#fig-c}']
        check('a figure inside a callout is counted: the emitter renders one',
              xref.resolve(profile.load('en'), body)['fig-c']['number'] == 2)
        html = render(body, root, svgs=[SVG])
        check('and the page agrees',
              '<figcaption>Figure 2. The diagram.' in html)

        body = ['- an item', '  ![a](f.png)', '', '```mermaid', 'graph TD',
                'A-->B', '```', '', ': The diagram. {#fig-d}']
        check('an image indented under a list is list content, not a float',
              xref.resolve(profile.load('en'), body)['fig-d']['number'] == 1
              and xref.floats(body) == [{'kind': 'fig', 'slot': 8}])

        body = ['# TITLE', '', '![a](f.png)', '', '---', '', '## Body', '',
                '```mermaid', 'graph TD', 'A-->B', '```', '',
                ': The diagram. {#fig-e}']
        check('a float in the head is rendered by nothing and numbers nothing',
              xref.resolve(profile.load('en'), body, head=5)['fig-e']['number'] == 1)

        print('a missing file leaves a visible gap, never a silent one')
        html = render(['![A field](gone.png)'], root)
        check('the gap says what is missing',
              'image not found: gone.png' in html)
        check('nothing pretends to be an image', '<img' not in html)

    print('lint refuses what the build cannot honour')
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'f.png').write_bytes(PNG)
        doc = document(root, ['![here](f.png)', '', ': A caption. {#fig-a}', '',
                              'Text about @fig-a.'])
        check('a resolved image with an attached caption is clean',
              lint.check_images(doc) == [] and lint.check_captions(doc) == [])

        # an included file is a fragment of one document, not a document, so
        # every reader of "the source" resolves against the source's directory.
        # What matters is that lint and the emitters agree on which directory
        # that is; they used to be able to differ silently.
        (root / 'parts').mkdir()
        (root / 'parts' / 'ch.md').write_text('![here](f.png)\n', encoding='utf-8')
        doc = document(root, ['Body.'], 'inc.md')
        doc['include_paths'] = (str(root / 'parts' / 'ch.md'),)
        check('an image in an included file resolves from the document',
              lint.check_images(doc) == [])

        doc = document(root, ['![gone](missing.png)'], 'b.md')
        check('a missing file blocks',
              rules(lint.check_images(doc)) == {('missing-image', 'block')})

        doc = document(root, ['![remote](https://example.com/f.png)'], 'c.md')
        check('a remote image blocks',
              rules(lint.check_images(doc)) == {('remote-image', 'block')})

        doc = document(root, ['Just prose here.', '', ': A caption. {#fig-a}', '',
                              'And @fig-a points at it.'], 'd.md')
        found = lint.check_captions(doc)
        check('a caption under prose blocks',
              rules(found) == {('stray-caption', 'block')})
        check('and it is reported on its own line',
              found and found[0]['line'] == 3)
        check('the older checks cannot see it: the label exists and is used',
              lint.check_references(doc, profile.load('en')) == [])

    print('what may carry a caption')
    check('not a caption under a list that contains an image',
          xref.attached_captions(['- an item', '  ![a](x.png)', '',
                                  ': A caption {#fig-a}']) == set())
    slots = xref.attached_captions(
        ['```mermaid', 'graph TD', '```', '', ': one {#fig-a}', '',
         '| a | b |', '|---|---|', '| 1 | 2 |', ': two {#tbl-a}', '',
         '![x](y.png)', '', ': three {#fig-b}', '',
         '- a list item', '', ': four {#fig-c}'])
    check('a diagram may', 4 in slots)
    check('a table may', 9 in slots)
    check('an image may', 13 in slots)
    check('a list may not', 17 not in slots)

    print('coverage does not ask a picture to be text')
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'f.png').write_bytes(PNG)
        src = root / 'doc.md'
        src.write_text('![A photograph of a salinity-affected field](f.png)\n',
                       encoding='utf-8')
        html = root / 'doc.html'
        html.write_text('<html><body><figure><img src="data:image/png;base64,x">'
                        '</figure></body></html>', encoding='utf-8')
        check('an image line is not reported as missing content',
              verify.coverage(str(html), str(src)) == [])

    print('the print edition sets the image, rather than deleting it')
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'f.png').write_bytes(PNG)
        typst.SRC['dir'] = root
        typst.PLATES.clear()
        typst.FLOATS.clear()
        body = ['![A field](f.png)', '', ': The photograph. {#fig-a}']
        typst.XREF.clear()
        typst.XREF.update(xref.resolve(profile.load('en'), body))
        out = typst.convert(body, {}, [], 'Figure %d')
        check('a block image becomes a Typst figure', '#figure(image("plate-0.png"' in out)
        check('its caption travels with it', 'The photograph' in out)
        check('and the file is registered to be copied beside the source',
              [n for n, _ in typst.PLATES] == ['plate-0.png'])

        (root / 'chart.svg').write_text('<svg viewBox="0 0 10 10"/>', encoding='utf-8')
        typst.PLATES.clear()
        typst.FLOATS.clear()
        typst.convert(['![A chart](chart.svg)'], {}, [], 'Figure %d')
        check('an SVG is rasterised for print, not placed as vector',
              [n for n, _ in typst.PLATES] == ['plate-0-0.png'])

        typst.PLATES.clear()
        typst.FLOATS.clear()
        out = typst.convert(['Prose with ![a mark](f.png) in it.'], {}, [], 'Figure %d')
        check('an inline image is set in the line', '#box(image("plate-0.png"' in out)
        check('and is not counted as a figure', typst.FLOATS == [])

        typst.PLATES.clear()
        typst.FLOATS.clear()
        out = typst.convert(['![A field](gone.png)'], {}, [], 'Figure %d')
        check('a missing file does not reach the Typst source',
              'image(' not in out and 'image not found' in out)
        check('and a missing inline one is not silently dropped',
              'image not found' in typst.inline('Prose ![a](gone.png) here.', {}))

    print('an image named in the front matter is copied too')
    if not require.found('typst'):
        print('  %-58s skip (typst is not installed)' % 'the print build finds it')
    else:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            (root / 'f.png').write_bytes(PNG)
            # the front matter renders last, after the block pass, so an image
            # named only there registers a plate after the body is converted
            (root / 'doc.md').write_text(FRONT_DOC, encoding='utf-8')
            out = root / 'doc.pdf'
            try:
                typst.build(str(root / 'doc.md'), str(out), profile.load('en'))
                built, why = out.is_file(), ''
            except RuntimeError as e:
                built, why = False, str(e)[:60]
            check('the print build finds it' + (' (%s)' % why if why else ''), built)

    print('the Word edition places the picture, and not over a diagram')
    calls = []
    real = typst.browser.run
    typst.browser.run = lambda args, **kw: calls.append(
        next(a for a in args if a.startswith('--screenshot=')))
    try:
        with tempfile.TemporaryDirectory() as tmp:
            typst.rasterise(['<svg viewBox="0 0 10 10"/>'], Path(tmp), prefix='mark')
            check('the project mark rasterises under its own name, not fig-0.png',
                  calls and calls[0].endswith('mark-0.png'))
    finally:
        typst.browser.run = real

    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / 'f.png').write_bytes(PNG)
        docx_mod.SRC['dir'] = root
        docx_mod.WORK['dir'] = None
        docx_mod.FLOATS.clear()
        doc = Document()
        figures = []
        docx_mod.convert(doc, ['![A field](f.png)', '', ': The photograph. {#fig-a}'],
                         figures, 'Figure %d', {}, {},
                         table=xref.resolve(profile.load('en'),
                                            ['![A field](f.png)', '',
                                             ': The photograph. {#fig-a}']))
        check('a picture is placed in the .docx',
              any('image' in r.reltype for r in doc.part.rels.values()))
        check('with its caption under it',
              any('The photograph' in p.text for p in doc.paragraphs))
        check('and it is not counted as a diagram', figures == [])

        docx_mod.FLOATS.clear()
        docx_mod.RASTERS.clear()
        doc = Document()
        para = doc.add_paragraph()
        docx_mod._runs(para, 'Prose with ![a mark](f.png) in it.')
        check('an inline image is placed, not written as its own markup',
              '![a mark]' not in para.text and '!a mark' not in para.text)
        check('and the picture is really there',
              any('image' in r.reltype for r in doc.part.rels.values()))

        para = doc.add_paragraph()
        docx_mod._runs(para, 'Documented as `![alt](f.png)` in the guide.')
        check('an image in a code span is left as code', '![alt](f.png)' in para.text)

        doc = Document()
        docx_mod.FLOATS.clear()
        docx_mod.convert(doc, ['![A field](gone.png)'], [], 'Figure %d', {}, {})
        check('a missing file leaves a visible note in Word too',
              any('image not found: gone.png' in p.text for p in doc.paragraphs))

    print()
    if failures:
        print('%d check(s) failed' % len(failures))
        return 1
    print('all checks passed')
    return 0


if __name__ == '__main__':
    sys.exit(main())
